"""
RAG (Retrieval Augmented Generation) functions for AtlasMind
Supports both video (YouTube) and PDF inputs; same pipeline for notes, quiz, summary.
"""

import gradio as gr
from models import current_video
from vector_db import semantic_search, chunk_text
from llm import ask_groq
from config import TRANSCRIPT_PREVIEW_LENGTH


def _process_content_text(content_id: str, transcript: str, source_label: str, source_type: str) -> str:
    """Store text in vector DB, set state, and generate summary. Shared by video and PDF."""
    from vector_db import store_in_vector_db

    current_video.transcript = transcript
    current_video.video_id = content_id
    current_video.source_type = source_type
    current_video.collection = store_in_vector_db(content_id, transcript)

    print("Generating AI summary...")
    prompt = f"""You are AtlasMind, an AI learning companion.

Analyze this content and provide:

## 📝 Summary
A concise 3-4 sentence overview.

## 🔑 Key Concepts
5-7 most important points as bullets.

## 💡 Takeaways
3-5 actionable insights.

Content: {transcript[:TRANSCRIPT_PREVIEW_LENGTH]}"""

    summary = ask_groq(prompt)
    print("Summary generated!")
    return f"""**{source_label} Processed Successfully!**

---

{summary}"""


def process_content(video_url: str = None, pdf_file=None, progress=gr.Progress()) -> str:
    """
    Process either a YouTube video or an uploaded PDF: extract text, store in DB, generate summary.
    Exactly one of video_url or pdf_file must be provided.

    Returns:
        Summary and processing status
    """
    from youtube import fetch_transcript_ytdlp
    from pdf import extract_text_from_pdf

    progress(0, desc="Fetching content...")
    video_url = (video_url or "").strip()
    # Gradio file upload can be single path or list of paths
    pdf_path = None
    if pdf_file is not None:
        if isinstance(pdf_file, list) and len(pdf_file) > 0:
            pdf_path = pdf_file[0] if isinstance(pdf_file[0], str) else getattr(pdf_file[0], "name", None)
        elif isinstance(pdf_file, str):
            pdf_path = pdf_file

    if video_url and pdf_path:
        return "Please provide either a video link OR a PDF file, not both."
    if not video_url and not pdf_path:
        return "Please enter a YouTube URL or upload a PDF file."

    progress(0.2, desc="Extracting content...")
    if video_url:
        print("\n" + "="*60)
        print("PROCESSING VIDEO (using yt-dlp)")
        print("="*60)
        result = fetch_transcript_ytdlp(video_url)
        if not result["success"]:
            return result["error"]
        content_id = result["video_id"]
        transcript = result["transcript"]
        source_label = "Video"
    else:
        print("\n" + "="*60)
        print("PROCESSING PDF")
        print("="*60)
        result = extract_text_from_pdf(pdf_path)
        if not result["success"]:
            return result["error"]
        content_id = result["content_id"]
        transcript = result["transcript"]
        source_label = "PDF"

    progress(0.6, desc="Generating summary...")
    print("="*60 + "\n")
    result_text = _process_content_text(content_id, transcript, source_label, "video" if video_url else "pdf")
    progress(1.0, desc="Done!")
    return result_text


def answer_question(question: str) -> str:
    """
    Answer questions using RAG
    
    Args:
        question: User's question
    
    Returns:
        Answer based on video context
    """
    if not current_video.transcript:
        return "Please process a video or PDF first!"
    
    if not question or not question.strip():
        return "Please enter a question!"
    
    print(f"\nQuestion: {question}")
    
    context = semantic_search(question, current_video.collection)
    if not context:
        context = current_video.transcript[:3000]
    
    print(f"Found relevant context ({len(context)} chars)")
    
    prompt = f"""Based on this content (lecture or document), answer the question clearly and concisely.

Question: {question}

Provide a helpful answer."""
    
    print(f"Generating answer...")
    answer = ask_groq(prompt, context)
    print(f"Answer ready!\n")
    
    return answer


def generate_notes() -> tuple:
    """
    Generate detailed study notes from content (~1.5 pages): paragraphs and bullets that explain the material.

    Returns:
        Tuple of (notes_markdown, file_path) for display and download. file_path is None on error.
    """
    import tempfile
    import os

    if not current_video.transcript:
        return ("Please process a video or PDF first!", None)

    print("\nGenerating detailed notes...")

    prompt = f"""Create DETAILED study notes from this content (lecture or document). Aim for about 1.5 pages of a Word document — substantial enough for a candidate to learn from.

Rules:
- Use CONTENT-SPECIFIC section headings. Extract actual topics from the material and use them as titles (e.g. "How Backpropagation Works", "Causes of Market Failure") — NOT generic headings like "Main Topics" or "Key Concepts".
- Mix short paragraphs and bullet points. Use paragraphs to explain ideas clearly; use bullets for lists, steps, or key takeaways.
- Write to EXPLAIN: the candidate should understand the material, not just see keywords. Define terms, give context, and connect ideas.
- Cover the main concepts in depth. Include definitions, explanations, and where relevant, examples or implications.

Content:
{current_video.transcript[:6000]}"""

    notes = ask_groq(prompt)
    print("Notes generated!\n")

    file_path = None
    try:
        file_path = _notes_to_docx(notes)
    except Exception as e:
        print(f"Could not save notes to file: {e}")

    return (notes, file_path)


def _notes_to_docx(markdown_text: str):
    """Convert markdown notes to a .docx file; return path to the file."""
    import tempfile
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

    doc = Document()
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # Heading (## or ###)
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)
        i += 1

    fd, path = tempfile.mkstemp(suffix=".docx", prefix="atlasmind_notes_")
    import os
    os.close(fd)
    doc.save(path)
    return path
